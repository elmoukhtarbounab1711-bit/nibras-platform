"""
خدمات مولّد الوثائق (المرحلة 4) — قرار D-022.

توليد وثائق قانونية من قوالب بيانات (field_schema + body_template بنمط
Jinja2) عبر سؤال موجَّه: تحقق صارم من الإجابات ثم استبدال في هيكل ثابت
وحفظ لكل مستخدم متدرج (version) وتصدير PDF/DOCX في الذاكرة بلا ملفات
على القرص. القوالب بيانات لا كود (المواصفة التقنية §6) — تُضاف أنواع
وثائق جديدة عبر السجل لا إعادة نشر. التصدير: DOCX بـ python-docx
(ضبط RTL)، PDF بـ reportlab + arabic-reshaper + python-bidi مع خط عربي
يُحلَّل من مسارات شائعة (D-022 — استُبعد WeasyPrint لتعذره على ويندوز).
"""
import io
import json
import os
import re
from datetime import date

from jinja2 import Environment, StrictUndefined
from jinja2.exceptions import UndefinedError

from . import config
from .database import db_session

_FIELD_TYPES = {"text", "textarea", "number", "date", "select", "boolean"}

# قوالب العقد والوثائق — في seed_templates.py (بيانات لا كود). أُعيدت صياغتها
# بشكلياتها الكاملة: تعيين الأطراف ببطاقة التعريف الوطنية، تمهيد، بنود
# مرقّمة بمرجعية قانونية، وبياض توقيع — لا نماذج تجريبية مختصرة. إدارة
# المحتوى (FR-12.1) مؤجَّلة كوحدة أدمن — النمط ذاته في D-021.
from .seed_templates import _SEED_TEMPLATES

# مسارات خطوط عربية شائعة (ويندوز ثم لينكس) — تُحلَّل عند أول تصدير PDF
# ويُتجاوز المسار عبر NIBRAS_PDF_FONT (config.PDF_FONT_PATH)
_CANDIDATE_FONTS = [
    ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
    ("/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf",
     "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Bold.ttf"),
    ("/usr/share/fonts/opentype/noto/NotoNaskhArabic-Regular.ttf", None),
    ("/usr/share/fonts/truetype/amiri/Amiri-Regular.ttf", None),
]

_FONT_REGISTERED = {"regular": None, "bold": None}


class DocumentError(Exception):
    def __init__(self, message: str, status_code: int = 400):
        super().__init__(message)
        self.message = message
        self.status_code = status_code


def _ensure_description_column(conn):
    """ترحيل خفيف لقواعد بيانات أنشئت قبل وصف القوالب (idempotent)."""
    cols = [row[1] for row in conn.execute("PRAGMA table_info(document_templates)")]
    if "description" not in cols:
        conn.execute(
            "ALTER TABLE document_templates ADD COLUMN description TEXT NOT NULL DEFAULT ''"
        )


def ensure_defaults():
    """بذر القوالب النموذجية (idempotent) — تُستدعى من init_db."""
    with db_session() as conn:
        _ensure_description_column(conn)
        for tmpl in _SEED_TEMPLATES:
            existing = conn.execute(
                "SELECT id FROM document_templates WHERE slug = ?", (tmpl["slug"],)
            ).fetchone()
            if existing:
                continue
            conn.execute(
                "INSERT INTO document_templates "
                "(slug, name, category, description, field_schema, body_template, created_at) "
                "VALUES (?, ?, ?, ?, ?, ?, datetime('now'))",
                (
                    tmpl["slug"],
                    tmpl["name"],
                    tmpl["category"],
                    tmpl.get("description", ""),
                    json.dumps(tmpl["field_schema"], ensure_ascii=False),
                    tmpl["body_template"],
                ),
            )


def _parse_schema(field_schema: str) -> list:
    try:
        fields = json.loads(field_schema)
    except (TypeError, ValueError):
        raise DocumentError("مخطط قالب الوثيقة غير صالح.", 500)
    if not isinstance(fields, list):
        raise DocumentError("مخطط قالب الوثيقة غير صالح.", 500)
    return fields


def list_templates(category: str | None = None):
    query = """
        SELECT id, slug, name, category, description
        FROM document_templates
    """
    params = []
    if category:
        query += " WHERE category = ?"
        params.append(category)
    query += " ORDER BY category, name"
    with db_session() as conn:
        rows = conn.execute(query, params).fetchall()
        return [dict(r) for r in rows]


def get_template(slug: str):
    with db_session() as conn:
        row = conn.execute(
            "SELECT id, slug, name, category, description, field_schema, body_template "
            "FROM document_templates WHERE slug = ?",
            (slug,),
        ).fetchone()
        if not row:
            return None
        tmpl = dict(row)
        tmpl["fields"] = _parse_schema(tmpl.pop("field_schema"))
        return tmpl


def _get_template_by_ident(ident):
    """يحل القالب بالمعرّف الرقمي (template_id) أو بالرابط (template_slug)."""
    with db_session() as conn:
        if isinstance(ident, int):
            row = conn.execute(
                "SELECT id, slug, name, category, description, field_schema, body_template "
                "FROM document_templates WHERE id = ?",
                (ident,),
            ).fetchone()
        else:
            row = conn.execute(
                "SELECT id, slug, name, category, description, field_schema, body_template "
                "FROM document_templates WHERE slug = ?",
                (str(ident),),
            ).fetchone()
        if not row:
            return None
        tmpl = dict(row)
        tmpl["fields"] = _parse_schema(tmpl.pop("field_schema"))
        return tmpl


def _validate_answers(fields: list, answers: dict) -> dict:
    """تحقق صارم من كل إجابة حسب مخطط الحقل — يرفض بأول خطأ (رسالة عربية)."""
    cleaned = {}
    for field in fields:
        name = field["name"]
        label = field.get("label") or name
        ftype = field.get("type", "text")
        if ftype not in _FIELD_TYPES:
            raise DocumentError(f"نوع حقل غير معروف في القالب: {name}", 500)
        raw = answers.get(name)
        if raw is None or (isinstance(raw, str) and not raw.strip()):
            if field.get("required"):
                raise DocumentError(f"الحقل «{label}» مطلوب.", 400)
            # قيم افتراضية للحقول الاختيارية حتى تمر بـ StrictUndefined داخل
            # {% if %} (فراغ/None → falsy فلا يُطبع شيء) — قيد D-022
            cleaned[name] = None if ftype in ("number", "boolean") else ""
            continue
        if ftype in ("text", "textarea"):
            cleaned[name] = str(raw).strip()
        elif ftype == "number":
            try:
                value = float(str(raw).replace(",", "").strip())
            except (TypeError, ValueError):
                raise DocumentError(f"الحقل «{label}» يجب أن يكون رقمًا.", 400)
            value = int(value) if value.is_integer() else value
            if "min" in field and value < field["min"]:
                raise DocumentError(f"الحقل «{label}» لا يمكن أن يقل عن {field['min']}.", 400)
            if "max" in field and value > field["max"]:
                raise DocumentError(f"الحقل «{label}» لا يمكن أن يزيد عن {field['max']}.", 400)
            cleaned[name] = value
        elif ftype == "date":
            text = str(raw).strip()
            try:
                date.fromisoformat(text)
            except ValueError:
                raise DocumentError(
                    f"الحقل «{label}» يجب أن يكون تاريخًا بصيغة YYYY-MM-DD.", 400
                )
            cleaned[name] = text
        elif ftype == "select":
            options = field.get("options") or []
            if raw not in options:
                raise DocumentError(f"القيمة المقدمة للحقل «{label}» غير مقبولة.", 400)
            cleaned[name] = raw
        elif ftype == "boolean":
            if isinstance(raw, bool):
                cleaned[name] = raw
            elif isinstance(raw, str) and raw.strip().lower() in ("true", "false"):
                cleaned[name] = raw.strip().lower() == "true"
            else:
                raise DocumentError(f"الحقل «{label}» يجب أن يكون قيمة منطقية.", 400)
    return cleaned


def _render_body(body_template: str, cleaned_answers: dict) -> str:
    """يستبدل الإجابات في الهيكل الثابت بنمط Jinja2 مع StrictUndefined —
    أي متغير مفقود خطأ صريح لا فراغ صامت (قرار D-022)."""
    try:
        env = Environment(undefined=StrictUndefined)
        return env.from_string(body_template).render(**cleaned_answers)
    except UndefinedError as exc:
        raise DocumentError("تعذر توليد الوثيقة من الإجابات المقدمة.", 400) from exc


# علامات تنسيق خفيفة تُضاف لنص الوثيقة المولَّد ويُفسِّرها مخرجا PDF/DOCX:
_M_TITLE = "#T#"   # عنوان مركزي عريض
_M_HEADING = "#H#" # عنوان بند (يمين، عريض)
_M_CENTER = "#C#"  # سطر مركزي (البسملة)
# تنسيق عريض سطريًا عبر **...** حول أسماء الأطراف

_PARTY_RE = re.compile(r"(السيد/السيدة|السيدة/السيد)\s+([^،\n]+?)(\s*[،؛]|$)")
_LABEL_RE = re.compile(r"^([^:]+?):\s+______")


def _format_doc_text(doc_text: str) -> str:
    """يرفع نص الوثيقة المولَّد إلى الشكل الاحترافي: البسملة والعنوان
    مركزيان، «المادة/الطرف/تمهيد» عناوين، وأسماء الأطراف بخط عريض."""
    lines = doc_text.split("\n")
    non_empty = [i for i, ln in enumerate(lines) if ln.strip()]
    basmala_idx = None
    title_idx = None
    if non_empty and lines[non_empty[0]].strip().startswith("بسم الله"):
        basmala_idx = non_empty[0]
        if len(non_empty) > 1:
            title_idx = non_empty[1]
    out = []
    for i, ln in enumerate(lines):
        s = ln.strip()
        if not s:
            out.append("")
            continue
        if i == basmala_idx:
            out.append(_M_CENTER + s)
            continue
        if i == title_idx:
            out.append(_M_TITLE + s)
            continue
        if s.startswith(("المادة ", "المادة —", "الطرف الأول:", "الطرف الثاني:")) \
                or s in ("تمهيد:", "الأطراف:"):
            out.append(_M_HEADING + s)
            continue
        s = _LABEL_RE.sub(r"**\1:** ______", s)
        s = _PARTY_RE.sub(r"\1 **\2**\3", s)
        out.append(s)
    return "\n".join(out)


def _extract_special_conditions(answers: dict) -> str:
    """حقل «شروط خاصة» اختياري يُقبل مع أي قالب (ليس ضمن المخطط) —
    يُدمج لاحقًا كبند رقمي في الوثيقة. يتيح إضافة شروط مخصصة دون حذفها
    من مخزون الإجابات حتى تظهر عند إعادة التوليد."""
    raw = (answers or {}).get("special_conditions")
    if raw is None:
        return ""
    if isinstance(raw, str):
        text = raw.strip()
    else:
        text = str(raw).strip()
    if not text:
        return ""
    return text


def _insert_special_conditions(doc_text: str, conditions: str) -> str:
    """يُدرج «شروط خاصة» كبند رقمي قبل بياض التوقيع (أو قبل التنبيه إن لم
    يوجد بياض) — يحافظ على الشكليات: بنود قبل التذييل."""
    clause = "\n\nالمادة — شروط خاصة:\n" + conditions.strip() + "\n"
    idx = doc_text.find("______")
    if idx != -1:
        return doc_text[:idx] + clause + "\n" + doc_text[idx:]
    idx = doc_text.find("تنبيه:")
    if idx != -1:
        return doc_text[:idx] + clause + "\n" + doc_text[idx:]
    return doc_text.rstrip() + "\n" + clause


def _document_payload(row) -> dict:
    row = dict(row)
    payload = {
        "id": row["id"],
        "template_id": row["template_id"],
        "template_slug": row["template_slug"],
        "template_name": row["template_name"],
        "version": row["version"],
        "doc_text": row["doc_text"],
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }
    if "answers_json" in row:
        try:
            payload["answers"] = json.loads(row["answers_json"])
        except (TypeError, ValueError):
            payload["answers"] = {}
    return payload


def generate_document(user_id: int | None, template_id, answers: dict,
                      format_: str = "text"):
    """يولّد وثيقة من قالب — عديم الحالة للزوار (منصة عامة).

    عندما يكون user_id هو None (زائر مجهول) لا تُخزَّن الوثيقة نهائيًا
    (لا سجلّ وثائق شخصي — الخصوصية بالتصميم): يُعاد النص المولَّد فقط،
    أو ملف PDF/DOCX فورًا إن طُلب format_ (توليد في الذاكرة بلا قرص).
    user_id غير None يستخدمه الاختبار الداخلي/الأدمن فقط (يحفظ متدرجًا).
    """
    tmpl = _get_template_by_ident(template_id)
    if not tmpl:
        raise DocumentError("القالب غير موجود.", 404)
    cleaned = _validate_answers(tmpl["fields"], answers or {})
    doc_text = _render_body(tmpl["body_template"], cleaned)
    conditions = _extract_special_conditions(answers)
    if conditions:
        doc_text = _insert_special_conditions(doc_text, conditions)
        cleaned["special_conditions"] = conditions
    if user_id is None:
        base = {
            "id": None,
            "template_id": tmpl["id"],
            "template_slug": tmpl["slug"],
            "template_name": tmpl["name"],
            "version": 1,
            "doc_text": doc_text,
            "created_at": None,
            "updated_at": None,
            "answers": cleaned,
            "persisted": False,
        }
        if format_ in ("pdf", "docx"):
            doc = dict(base)
            if format_ == "docx":
                base["data"] = export_docx(doc)
            else:
                base["data"] = export_pdf(doc)
        return base
    with db_session() as conn:
        cur = conn.execute(
            "INSERT INTO generated_documents "
            "(user_id, template_id, answers_json, version, doc_text, created_at, updated_at) "
            "VALUES (?, ?, ?, 1, ?, datetime('now'), datetime('now'))",
            (user_id, tmpl["id"], json.dumps(cleaned, ensure_ascii=False), doc_text),
        )
        row = conn.execute(
            "SELECT g.id, g.template_id, t.slug AS template_slug, t.name AS template_name, "
            "       g.version, g.doc_text, g.answers_json, g.created_at, g.updated_at "
            "FROM generated_documents g JOIN document_templates t ON t.id = g.template_id "
            "WHERE g.id = ?",
            (cur.lastrowid,),
        ).fetchone()
    return _document_payload(row)


def get_user_documents(user_id: int):
    with db_session() as conn:
        rows = conn.execute(
            "SELECT g.id, g.template_id, t.slug AS template_slug, t.name AS template_name, "
            "       g.version, g.doc_text, g.answers_json, g.created_at, g.updated_at "
            "FROM generated_documents g JOIN document_templates t ON t.id = g.template_id "
            "WHERE g.user_id = ? ORDER BY g.created_at DESC, g.id DESC",
            (user_id,),
        ).fetchall()
        return [_document_payload(r) for r in rows]


def get_document(user_id: int, doc_id: int):
    """يعيد وثيقة المستخدم مع تحقق الملكية — الغير يرى 404 (owner-only)."""
    with db_session() as conn:
        row = conn.execute(
            "SELECT g.id, g.template_id, t.slug AS template_slug, t.name AS template_name, "
            "       g.version, g.doc_text, g.answers_json, g.created_at, g.updated_at "
            "FROM generated_documents g JOIN document_templates t ON t.id = g.template_id "
            "WHERE g.id = ? AND g.user_id = ?",
            (doc_id, user_id),
        ).fetchone()
        if not row:
            raise DocumentError("الوثيقة غير موجودة.", 404)
        return _document_payload(row)


def regenerate_document(user_id: int, doc_id: int, answers: dict):
    """FR-5.3 متدرج عند التعديل: نسخة +1 مع تحديث doc_text (قرار D-022)."""
    doc = get_document(user_id, doc_id)
    tmpl = _get_template_by_ident(doc["template_id"])
    cleaned = _validate_answers(tmpl["fields"], answers or {})
    doc_text = _render_body(tmpl["body_template"], cleaned)
    conditions = _extract_special_conditions(answers)
    if conditions:
        doc_text = _insert_special_conditions(doc_text, conditions)
        cleaned["special_conditions"] = conditions
    with db_session() as conn:
        conn.execute(
            "UPDATE generated_documents SET answers_json = ?, version = version + 1, "
            "doc_text = ?, updated_at = datetime('now') WHERE id = ?",
            (json.dumps(cleaned, ensure_ascii=False), doc_text, doc_id),
        )
        row = conn.execute(
            "SELECT g.id, g.template_id, t.slug AS template_slug, t.name AS template_name, "
            "       g.version, g.doc_text, g.answers_json, g.created_at, g.updated_at "
            "FROM generated_documents g JOIN document_templates t ON t.id = g.template_id "
            "WHERE g.id = ?",
            (doc_id,),
        ).fetchone()
    return _document_payload(row)


def _split_segments(text: str):
    """يقسّم سطرًا على وسم **...** إلى (نص، عريض) مع الحفاظ على ترتيب السطر."""
    parts = re.split(r"\*\*([^*]+)\*\*", text)
    return [(part, i % 2 == 1) for i, part in enumerate(parts) if part]


def _docx_style_marker(text: str):
    """يكشف وسم التنسيق في بداية السطر ويعيد (الوسم، بقية النص) أو (None, text)."""
    for marker in (_M_TITLE, _M_HEADING, _M_CENTER):
        if text.startswith(marker):
            return marker, text[len(marker):].strip()
    return None, text


def export_docx(doc: dict) -> bytes:
    """DOCX عربي بضبط اتجاه الفقرات RTL — توليد في الذاكرة (استيراد مؤجَّل
    فلا اعتماد صلب إن لم تُثبَّت المكتبة، نمط Anthropic في D-021)."""
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError:
        raise DocumentError("مكتبة تصدير DOCX غير مثبتة.", 500)
    buffer = io.BytesIO()
    d = DocxDocument()
    d.styles["Normal"].font.name = "Arial"
    formatted = _format_doc_text(doc["doc_text"] or "")
    lines = formatted.split("\n")
    for line in lines:
        text = line.strip()
        if not text:
            continue
        marker, text = _docx_style_marker(text)
        p = d.add_paragraph()
        if marker in (_M_TITLE, _M_CENTER):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
        for segment, bold in _split_segments(text):
            run = p.add_run(segment)
            run.font.name = "Arial"
            if marker == _M_TITLE:
                run.font.size = Pt(18)
            elif marker == _M_HEADING:
                run.font.size = Pt(13)
            elif marker == _M_CENTER:
                run.font.size = Pt(14)
            if bold or marker in (_M_TITLE, _M_HEADING):
                run.bold = True
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:cs"), "Arial")
    d.save(buffer)
    return buffer.getvalue()


def _resolve_pdf_font():
    """يسجّل خطًا عربيًا لدى reportlab مرة واحدة ويعيد (regular, bold)."""
    if _FONT_REGISTERED["regular"] is not None:
        return _FONT_REGISTERED["regular"], _FONT_REGISTERED["bold"]
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise DocumentError("مكتبة تصدير PDF غير مثبتة.", 500)
    candidates = []
    if config.PDF_FONT_PATH:
        candidates.append((config.PDF_FONT_PATH, None))
    candidates.extend(_CANDIDATE_FONTS)
    for regular, bold in candidates:
        if not (regular and os.path.exists(regular)):
            continue
        pdfmetrics.registerFont(TTFont("Arabic", regular))
        bold_name = "Arabic"
        if bold and os.path.exists(bold):
            try:
                pdfmetrics.registerFont(TTFont("Arabic-Bold", bold))
                bold_name = "Arabic-Bold"
            except (TimeoutError, OSError, ValueError, TypeError, KeyError):
                bold_name = "Arabic"
        try:
            pdfmetrics.registerFontFamily(
                "Arabic", normal="Arabic", bold=bold_name, italic="Arabic", boldItalic=bold_name
            )
        except (TimeoutError, OSError, ValueError, TypeError, KeyError):
            pass
        _FONT_REGISTERED["regular"] = "Arabic"
        _FONT_REGISTERED["bold"] = bold_name
        return "Arabic", bold_name
    raise DocumentError("خط عربي للتصدير غير متوفر على هذا النظام.", 500)


def _pdf_html_segments(segments, display_func, escape):
    """يبني نص Paragraph بوسوم <b> حول المقاطع العريضة بعد إعادة تشكيلها."""
    html = []
    for seg, bold in segments:
        if not seg:
            continue
        rendered = display_func(seg)
        rendered = escape(rendered)
        html.append(f"<b>{rendered}</b>" if bold else rendered)
    return "".join(html)


def export_pdf(doc: dict) -> bytes:
    """PDF عربي عبر reportlab (بايثون خالص) — إعادة ترتيب النص بـ
    arabic_reshaper + python-bidi ثم بناء المستند في الذاكرة."""
    try:
        from xml.sax.saxutils import escape

        import arabic_reshaper
        from bidi.algorithm import get_display
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        raise DocumentError("مكتبات تصدير PDF غير مثبتة.", 500)
    regular, bold = _resolve_pdf_font()
    buffer = io.BytesIO()
    base = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=25 * mm,
        bottomMargin=25 * mm,
        title="نبراس — وثيقة",
    )
    styles = {
        "title": ParagraphStyle(
            "title", fontName=bold, fontSize=16, leading=22, alignment=TA_CENTER, spaceAfter=14
        ),
        "heading": ParagraphStyle(
            "heading", fontName=bold, fontSize=13, leading=20, alignment=TA_RIGHT,
            spaceBefore=8, spaceAfter=4
        ),
        "center": ParagraphStyle(
            "center", fontName=regular, fontSize=13, leading=20, alignment=TA_CENTER, spaceAfter=8
        ),
        "body": ParagraphStyle(
            "body", fontName=regular, fontSize=12, leading=20, alignment=TA_RIGHT
        ),
    }
    story = []
    formatted = _format_doc_text(doc["doc_text"] or "")
    lines = formatted.split("\n")
    first = True
    for line in lines:
        text = line.strip()
        if not text:
            if not first:
                story.append(Spacer(1, 6))
            continue
        marker, text = _docx_style_marker(text)
        segments = _split_segments(text)
        html = _pdf_html_segments(
            segments, lambda seg: get_display(arabic_reshaper.reshape(seg)), escape
        )
        if marker == _M_TITLE:
            style = styles["title"]
        elif marker == _M_HEADING:
            style = styles["heading"]
        elif marker == _M_CENTER:
            style = styles["center"]
        else:
            style = styles["body"]
        story.append(Paragraph(html, style))
        first = False
    base.build(story)
    return buffer.getvalue()
